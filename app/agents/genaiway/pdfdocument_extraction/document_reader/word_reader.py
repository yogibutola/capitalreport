from docx import Document


class WordReader:

    def read_text_from_worddoc(self, file_path: str) -> list[str]:
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                text_parts.append(para.text.strip())

        table_parts = []
        for table_index, table in enumerate(doc.tables):
            table_parts.append(f"\n[Table {table_index + 1}]")
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_parts.append(" | ".join(row_data))

        # combined_text = "\n".join(text_parts + table_parts)

        return text_parts
