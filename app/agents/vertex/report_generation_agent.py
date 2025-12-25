import json

from dotenv import load_dotenv
from google.cloud import aiplatform
from pymongo.synchronous.collection import Collection
from vertexai.generative_models import (
    GenerationConfig,
    GenerativeModel,
    Part
)
from docx import Document

import base64
import re

from app.store.gcp_file_store import GCPStore
from app.store.mongo_db_store import MongoDBStore

load_dotenv()
GENERATION_MODEL = "gemini-2.5-flash"

import logging

logging.basicConfig(
    level=logging.INFO,  # Only output messages at INFO level and above
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)


class ReportGenerationAgent:
    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def generate_report(self, query: str):
        """
        Retrieves relevant documents and uses Gemini to generate a report.
        """

        # Format the retrieved context for the prompt
        # context = "\n---\n".join(context_blocks)
        file_part = Part.from_uri(
            uri="gs://capitalreport_file_storage/WTG- 24 Audit Report.pdf",
            mime_type="application/pdf"
        )

        # self.logger.debug(f"Context: {context}")
        # 2. Augmented Generation: Craft the RAG prompt
        prompt = f"""
        You are an expert report generation assistant for the provided document.
        Generate a roll-forward of the attached audit report for the fiscal years ending December 31, 2025, and 2024. Please perform the following specific actions:

        1. **Shift Dates:** Update all header dates, report dates, and fiscal year references throughout the document to reflect **2025 (Current Year)** and **2024 (Prior Year)**.
        2. **Remove Oldest Data:** Delete all financial data and columns referring to the year **2023**.
        3. **Move Prior Data:** Move the existing data from the **2024** columns into the **2024 (Prior Year)** columns.
        4. **Clear Current Data:** Create empty placeholders (e.g., `[INSERT]`) in the **2025 (Current Year)** columns for new data entry.
        5. **Output Format:**
            Output STRICT JSON with:
               - title
               - sections[]
               - tables[] (rows + columns)
               - placeholders as "[INSERT]"
               - no markdown
               - no base64
        """

        PROJECT_ID = "stable-smithy-270416"
        REGION = "us-central1"
        MODEL_NAME = "gemini-2.5-flash"
        # MODEL_NAME = "gemini-3-pro-preview"

        aiplatform.init(project=PROJECT_ID, location=REGION, request_metadata=[("grpc-timeout", "300S")])
        client = GenerativeModel(MODEL_NAME)
        config = GenerationConfig(
            temperature=0,  # Equivalent to temperature=0
            # system_instruction=INSTRUCTIONS,
            # tools=[GROUNDING_TOOL],         # Equivalent to tools=[grounding_tool]
            # response_mime_type="application/json", # Use this with response_schema
            # response_schema=RESPONSE_SCHEMA,      # Equivalent to response_schema=list[Investment]
        )

        try:
            response = client.generate_content(
                [file_part, prompt],
                generation_config=config,
            )

            data = json.loads(response.text)
            doc = Document()
            doc.add_heading(data["title"], level=1)

            for section in data["sections"]:
                doc.add_heading(section["heading"], level=2)
                doc.add_paragraph(section["text"])

            for table_data in data["tables"]:
                table = doc.add_table(
                    rows=len(table_data["rows"]),
                    cols=len(table_data["rows"][0]),
                )
                for r, row in enumerate(table_data["rows"]):
                    for c, cell in enumerate(row):
                        table.cell(r, c).text = cell

            doc.save("final.docx")

            #
            # raw_text = response.candidates[0].content.parts[0].text
            # cleaned = re.sub(r"```json|```", "", raw_text).strip()
            #
            # if cleaned.startswith('"') and cleaned.endswith('"'):
            #     cleaned = cleaned[1:-1]
            # cleaned = cleaned.replace("\n", "").strip()
            # docx_bytes = base64.b64decode(cleaned)
            #
            #
            # if not docx_bytes:
            #     raise RuntimeError("DOCX file was not returned by Gemini")
            # else:
            #     with open("debug.docx", "wb") as f:
            #         f.write(docx_bytes)

                # gcp_store = GCPStore()
                # gcp_store.upload_to_gcs(docx_bytes)

        except Exception as e:
            self.logger.error(f"An error occurred while generating the response: {e}")
            return None

        # self.logger.info(f"Response: {response_text}")
        return "Successfully generated the file."
